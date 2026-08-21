#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate DeepSeek-Harness-Handbook.docx (English, 14 chapters + Chinese-original appendices).

Reads docs/01-14 *.en.md (in numeric order) plus the Chinese appendix files
(appendix-glossary.md / appendix-packages.md / benchmark.md, appended as-is with
an English note), and produces a professionally typeset DOCX with:
  - cover page
  - static table of contents (no page numbers - LibreOffice headless cannot
    update a field-based TOC)
  - 14 chapters, each starting on a new page
  - Appendix section (Chinese original) with English labels
  - footer page numbers

Usage:
    python scripts/gen_en_pdf.py [repo_root] [out_docx]
"""

import datetime
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = sys.argv[1] if len(sys.argv) > 1 else "."
DOCS = os.path.join(REPO, "docs")
LANG = (sys.argv[3] if len(sys.argv) > 3 else os.environ.get("DSH_LANG", "en")).lower()
if LANG not in ("en", "zh"):
    raise SystemExit("lang must be 'en' or 'zh', got: %s" % LANG)
_DEFAULT_OUT = "DeepSeek-Harness-Handbook.docx" if LANG == "en" else "DeepSeek-Harness-Handbook-zh.docx"
OUT = sys.argv[2] if len(sys.argv) > 2 else os.path.join(REPO, _DEFAULT_OUT)

LATIN = "Calibri"
MONO = "Consolas"
CJK = "Microsoft YaHei"

NAVY = RGBColor(0x1F, 0x38, 0x64)
GRAY = RGBColor(0x59, 0x59, 0x59)
CODE_BG = "F2F2F2"
HEAD_BG = "DEEAF6"

BODY_SIZE = 11
CODE_SIZE = 9


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False, color=None, mono=False, cjk=True):
    run.font.name = MONO if mono else LATIN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), MONO if mono else LATIN)
    rfonts.set(qn("w:hAnsi"), MONO if mono else LATIN)
    if cjk:
        rfonts.set(qn("w:eastAsia"), CJK)


def shade_paragraph(p, fill):
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def h1_bottom_border(p):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1F3864")
    pbdr.append(bottom)
    ppr.append(pbdr)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*\n]+\*|\[[^\]]*\]\([^)\s]+\))")


def add_inline(p, text, size=BODY_SIZE, bold=False, italic=False, color=None, mono=False):
    """Add runs to paragraph `p`, parsing **bold**, *italic*, `code` and [text](url)."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            add_inline(p, part[2:-2], size, True, italic, color, mono)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = p.add_run(part[1:-1])
            set_run_font(r, size=size - 0.5, bold=bold, italic=italic, color=color, mono=True)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            add_inline(p, part[1:-1], size, bold, True, color, mono)
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]*)\]\(([^)\s]+)\)", part)
            if m:
                label, url = m.group(1), m.group(2)
                if url.startswith(("http://", "https://")):
                    add_inline(p, label, size, bold, italic, color, mono)
                    r = p.add_run("  (" + url + ")")
                    set_run_font(r, size=size - 2, bold=False, italic=True, color=GRAY)
                else:
                    add_inline(p, label, size, bold, italic, color, mono)
            else:
                r = p.add_run(part)
                set_run_font(r, size, bold, italic, color, mono)
        else:
            r = p.add_run(part)
            set_run_font(r, size, bold, italic, color, mono)


def body_paragraph(doc, text, size=BODY_SIZE, bold=False, italic=False, color=None,
                   space_after=8, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    add_inline(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def render_heading(doc, level, text, page_break=False):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    if page_break:
        fmt.page_break_before = True
    if level == 1:
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(14)
        fmt.keep_with_next = True
        h1_bottom_border(p)
        add_inline(p, text, size=16, bold=True, color=NAVY)
    elif level == 2:
        fmt.space_before = Pt(14)
        fmt.space_after = Pt(7)
        fmt.keep_with_next = True
        add_inline(p, text, size=13, bold=True, color=NAVY)
    elif level == 3:
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(5)
        fmt.keep_with_next = True
        add_inline(p, text, size=11.5, bold=True, italic=True)
    else:
        fmt.space_before = Pt(10)
        fmt.space_after = Pt(4)
        add_inline(p, text, size=10.5, bold=True)
    return p


def render_code(doc, lines):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0.3)
    fmt.right_indent = Cm(0.3)
    fmt.space_before = Pt(9)
    fmt.space_after = Pt(12)
    fmt.line_spacing = 1.1
    shade_paragraph(p, CODE_BG)
    for i, ln in enumerate(lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(ln if ln else " ")
        set_run_font(r, size=CODE_SIZE, mono=True)
    return p


def render_blockquote(doc, lines):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0.5)
    fmt.space_before = Pt(5)
    fmt.space_after = Pt(10)
    fmt.line_spacing = 1.4
    for i, ln in enumerate(lines):
        if i:
            p.add_run().add_break()
        add_inline(p, ln, size=10, italic=True, color=GRAY)
    return p


def render_list(doc, items, ordered=False):
    for idx, (marker, it) in enumerate(items):
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.left_indent = Cm(0.6)
        fmt.space_after = Pt(4)
        fmt.line_spacing = 1.5
        num = marker if ordered else "\u2022"
        r = p.add_run(num + "  ")
        set_run_font(r, size=BODY_SIZE, bold=False)
        add_inline(p, it, size=BODY_SIZE)


def render_table(doc, rows):
    ncols = max(len(r) for r in rows)
    header = rows[0]
    body = rows[2:] if len(rows) > 1 and re.match(r"^[\s:\-|]+$", "".join(rows[1])) else rows[1:]
    table = doc.add_table(rows=1 + len(body), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(header):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, c, size=10, bold=True)
        shade_cell(cell, HEAD_BG)
    for i, row in enumerate(body):
        for j in range(ncols):
            cell = table.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, row[j] if j < len(row) else "", size=10)
    # small spacer so following content does not butt against the table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    sp.paragraph_format.line_spacing = 1.0
    r = sp.add_run("")
    r.font.size = Pt(2)


SPLIT_ROW_RE = re.compile(r"^\s*\|(.*)\|\s*$")


def split_row(line):
    inner = SPLIT_ROW_RE.match(line).group(1)
    cells = []
    cur = ""
    esc = False
    for ch in inner:
        if esc:
            cur += ch
            esc = False
        elif ch == "\\":
            esc = True
        elif ch == "|":
            cells.append(cur.strip())
            cur = ""
        else:
            cur += ch
    cells.append(cur.strip())
    return cells


def clean_md_lines(raw_lines):
    """Remove nav header / footer lines that are for the web only."""
    lines = [l.rstrip("\n") for l in raw_lines]
    # drop leading nav line: [English](./x.en.md) | [中文](./x.md) · [← Back](../README.md)
    if lines and lines[0].lstrip().startswith("[English]"):
        lines = lines[1:]
    # drop '---' + '**Next**: ...' footer
    out = []
    for i, l in enumerate(lines):
        if l.strip() == "---" and i + 1 < len(lines) and lines[i + 1].strip().startswith("**Next**"):
            continue
        if l.strip().startswith("**Next**"):
            continue
        if re.match(r"^\s*<!--.*-->\s*$", l):
            continue
        out.append(l)
    return out


def render_markdown(doc, filepath, first_h1_as_heading=False, chapter_h1=True):
    """Render a markdown file into the docx. Returns the H1 title text if found."""
    with open(filepath, "r", encoding="utf-8") as f:
        lines = clean_md_lines(f.readlines())

    title = None
    i = 0
    n = len(lines)
    in_code = False
    code_buf = []
    quote_buf = []
    list_buf = []
    list_ordered = False
    table_buf = []

    def flush_code():
        if code_buf:
            render_code(doc, code_buf)
            code_buf.clear()

    def flush_quote():
        if quote_buf:
            render_blockquote(doc, quote_buf)
            quote_buf.clear()

    def flush_list():
        if list_buf:
            render_list(doc, list_buf, ordered=list_ordered)
            list_buf.clear()

    def flush_table():
        if table_buf:
            render_table(doc, table_buf)
            table_buf.clear()

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ---- code fence ----
        if stripped.startswith("```"):
            flush_code(); flush_quote(); flush_list(); flush_table()
            if in_code:
                render_code(doc, code_buf)
                code_buf.clear()
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            flush_quote(); flush_list(); flush_table()
            i += 1
            continue

        # ---- HTML helpers ----
        if stripped == "<details>" or stripped == "</details>":
            flush_quote(); flush_list(); flush_table()
            i += 1
            continue
        m = re.match(r"<summary>(.*)</summary>", stripped, re.S)
        if m:
            flush_quote(); flush_list(); flush_table()
            render_heading(doc, 4, m.group(1).strip())
            i += 1
            continue

        # ---- headings ----
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            flush_code(); flush_quote(); flush_list(); flush_table()
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1 and title is None:
                title = text
                render_heading(doc, 1, text, page_break=chapter_h1)
            else:
                render_heading(doc, level, text)
            i += 1
            continue

        # ---- tables ----
        if SPLIT_ROW_RE.match(stripped):
            flush_code(); flush_quote(); flush_list()
            table_buf.append(split_row(line))
            i += 1
            continue
        if table_buf:
            flush_table()

        # ---- horizontal rule ----
        if re.match(r"^\s*-{3,}\s*$", stripped) or re.match(r"^\s*\*{3,}\s*$", stripped):
            flush_quote(); flush_list()
            i += 1
            continue

        # ---- images ----
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if m:
            flush_code(); flush_quote(); flush_list()
            alt, path = m.group(1), m.group(2).strip()
            abs_path = os.path.join(os.path.dirname(filepath), path)
            if os.path.exists(abs_path):
                try:
                    pic_p = doc.add_paragraph()
                    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pic_p.paragraph_format.space_before = Pt(8)
                    pic_p.paragraph_format.space_after = Pt(8)
                    pic_p.add_run().add_picture(abs_path, width=Cm(13.5))
                except Exception as e:  # noqa: BLE001 - best-effort image embedding
                    print(f"  [warn] image embed failed: {abs_path}: {e}")
            i += 1
            continue

        # ---- blockquote ----
        if stripped.startswith(">"):
            flush_list(); flush_table()
            quote_buf.append(re.sub(r"^\s*>\s?", "", line))
            i += 1
            continue

        # ---- lists ----
        m = re.match(r"^(\s*)([-*])\s+(.*)$", line)
        if m:
            flush_quote(); flush_table()
            if list_buf and list_ordered:
                flush_list()
            list_ordered = False
            list_buf.append(("", m.group(3)))
            i += 1
            continue
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            flush_quote(); flush_table()
            if list_buf and not list_ordered:
                flush_list()
            list_ordered = True
            list_buf.append((m.group(2) + ".", m.group(3)))
            i += 1
            continue
        flush_list()

        # ---- plain paragraph ----
        flush_quote()
        body_paragraph(doc, stripped)
        i += 1

    flush_code(); flush_quote(); flush_list(); flush_table()
    return title


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_field(instr):
        r1 = p.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        r1._r.append(fld)
        r2 = p.add_run()
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = instr
        r2._r.append(it)
        r3 = p.add_run()
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        r3._r.append(fld2)

    r = p.add_run("Page ")
    set_run_font(r, size=9, color=GRAY)
    add_field("PAGE")
    r = p.add_run(" of ")
    set_run_font(r, size=9, color=GRAY)
    add_field("NUMPAGES")


def build_cover(doc, nchap, version, datestr):
    EN = (LANG == "en")
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_inline(p, "DeepSeek Harness Handbook" if EN else "DeepSeek Harness \u767d\u76ae\u4e66",
               size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    add_inline(p, "dsh-handbook \u00b7 Complete English Edition" if EN
               else "dsh-handbook \u00b7 \u4e2d\u6587\u5b8c\u6574\u7248", size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_inline(p,
               ("From zero to one with DeepSeek\u2019s open-source agent runtime \u2014 the "
                "beginner\u2019s encyclopedia: install, use, develop plugins, tune, secure, and budget.")
               if EN else
               ("\u4ece 0 \u5230 1 \u7528\u597d DeepSeek \u5f00\u6e90 Agent \u8fd0\u884c\u65f6\u2014\u2014"
                "\u5b89\u88c5\u3001\u4f7f\u7528\u3001\u5199\u63d2\u4ef6\u3001\u8c03\u4f18\u3001\u5b89\u5168\u4e0e\u6210\u672c\u3002"),
               size=12, italic=True, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    add_inline(p,
               ("%d Chapters + Appendix (Glossary / Packages / Benchmark)" % nchap) if EN
               else ("%d \u7ae0 + \u9644\u5f55 ABC\uff08\u672f\u8bed\u8868 / \u5b98\u65b9\u5305 / \u5b9e\u6d4b\u5bf9\u6bd4\uff09" % nchap),
               size=12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, "%s docs/ \u00b7 dsh %s \u00b7 %s"
               % ("Synced from" if EN else "\u540c\u6b65\u81ea", version, datestr),
               size=10, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, "https://github.com/Electricitysheep/dsh-handbook", size=10, color=GRAY)

    doc.add_page_break()


def build_toc(doc, chapters, appendices):
    render_heading(doc, 1, "Contents", page_break=False)
    for num, title in chapters:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.3
        add_inline(p, f"{num}  {title}", size=11)
    for label, cn_title in appendices:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.3
        add_inline(p, f"{label} \u2014 {cn_title}", size=11)
    doc.add_page_break()


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)   # A4
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # default paragraph font for empty paragraphs etc.
    style = doc.styles["Normal"]
    style.font.name = LATIN
    style.font.size = Pt(BODY_SIZE)
    rpr = style.element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), CJK)

    if LANG == "en":
        _pat = re.compile(r"^\d{2}-.*\.en\.md$")
    else:
        _pat = re.compile(r"^\d{2}-(?!.*\.en\.md).*\.md$")
    chapter_files = sorted(f for f in os.listdir(DOCS) if _pat.match(f))
    if not chapter_files:
        raise SystemExit("no chapter files found for lang=%s in %s" % (LANG, DOCS))
    appendix_files = [
        ("Appendix A", "appendix-glossary.md"),
        ("Appendix B", "appendix-packages.md"),
        ("Appendix C", "benchmark.md"),
    ]

    _version = os.environ.get("DSH_VERSION", "0.1.0-rc.8")
    _date = os.environ.get("DSH_DATE", datetime.date.today().isoformat())
    build_cover(doc, len(chapter_files), _version, _date)

    # collect titles for the TOC
    chapters = []
    for f in chapter_files:
        with open(os.path.join(DOCS, f), "r", encoding="utf-8") as fh:
            first_h1 = None
            for raw in clean_md_lines(fh.readlines()):
                m = re.match(r"^#\s+(.+)$", raw.strip())
                if m:
                    first_h1 = m.group(1).strip()
                    break
        num = f.split("-", 1)[0]
        chapters.append((num, first_h1 or f))

    appendix_cn_titles = ["\u672f\u8bed\u8868\u4e0e\u547d\u4ee4\u901f\u67e5",  # 术语表与命令速查
                          "\u5b98\u65b9\u5305\u901f\u67e5\u5927\u5168",      # 官方包速查大全
                          "\u540c\u6a21\u578b\u00d7\u4e0d\u540c Agent \u5b9e\u6d4b\u5bf9\u6bd4"]  # 同模型×不同Agent实测对比
    appendices = list(zip(appendix_files, appendix_cn_titles))
    build_toc(doc, chapters, [(label, cn) for (label, _), cn in appendices])

    # chapters
    for f in chapter_files:
        print(f"  rendering {f}")
        render_markdown(doc, os.path.join(DOCS, f), chapter_h1=True)

    # appendices (Chinese original, with English section note)
    render_heading(doc, 1, "Appendix · 附录（中文原文）" if LANG == "en" else "附录", page_break=True)
    if LANG == "en":
        _note = ("The appendices below are maintained in Chinese; they are included here in "
                 "their original form so the PDF stays complete. English readers can find the "
                 "equivalent concepts inside Chapters 1-%d." % len(chapter_files))
    else:
        _note = "\u4ee5\u4e0b\u9644\u5f55\u4e3a\u4e2d\u6587\u539f\u6587\uff0c\u968f\u6b63\u6587\u540c\u6b65\u7ef4\u62a4\u3002"
    body_paragraph(doc, _note, size=10, italic=True, color=GRAY, space_after=10)
    for label, fname in appendix_files:
        print(f"  rendering {fname} ({label})")
        render_heading(doc, 2, f"{label} \u00b7 {fname.replace('.md', '')} (Chinese original)")
        render_markdown(doc, os.path.join(DOCS, fname), chapter_h1=False)

    add_page_number_footer(doc)
    doc.save(OUT)
    print(f"SAVED: {OUT}")


if __name__ == "__main__":
    main()
